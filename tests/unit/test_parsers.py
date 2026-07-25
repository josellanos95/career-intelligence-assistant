import io

import pytest
from docx import Document as DocxDocument

from app.infra.parsers.docx_parser import DocxParser
from app.infra.parsers.factory import UnsupportedFileType, get_parser_for
from app.infra.parsers.pdf_parser import PdfParser
from app.infra.parsers.txt_parser import TxtParser


def test_txt_parser_decodes_utf8():
    assert TxtParser().parse("hello wörld".encode()) == "hello wörld"


def test_docx_parser_extracts_paragraphs():
    doc = DocxDocument()
    doc.add_paragraph("Hello world")
    doc.add_paragraph("Second paragraph")
    buffer = io.BytesIO()
    doc.save(buffer)

    text = DocxParser().parse(buffer.getvalue())

    assert "Hello world" in text
    assert "Second paragraph" in text


def test_pdf_parser_does_not_raise_on_valid_empty_pdf():
    import pymupdf

    doc = pymupdf.open()
    doc.new_page(width=200, height=200)
    buffer = io.BytesIO(doc.tobytes())

    text = PdfParser().parse(buffer.getvalue())

    assert isinstance(text, str)


def test_pdf_parser_preserves_line_breaks():
    import pymupdf

    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Your Mission")
    page.insert_text((72, 100), "We build things.")
    buffer = io.BytesIO(doc.tobytes())

    text = PdfParser().parse(buffer.getvalue())

    lines = [line.strip() for line in text.splitlines() if line.strip()]
    assert "Your Mission" in lines


def test_factory_returns_pdf_parser_for_pdf_extension():
    assert isinstance(get_parser_for("resume.pdf"), PdfParser)


def test_factory_returns_docx_parser_case_insensitively():
    assert isinstance(get_parser_for("resume.DOCX"), DocxParser)


def test_factory_returns_txt_parser_for_txt_extension():
    assert isinstance(get_parser_for("notes.txt"), TxtParser)


def test_factory_raises_for_unsupported_extension():
    with pytest.raises(UnsupportedFileType):
        get_parser_for("resume.pages")
